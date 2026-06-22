"""
Resume Writer Module
====================

Rewrites resume sections with improved wording based on job description
and analysis results. Supports PDF and DOCX export.

Copyright (c) 2025 SyazWak
Licensed under the MIT License - see LICENSE file for details.
"""

import os
import re
import json
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class SectionRewrite:
    """A single section rewrite with original and improved text."""
    section_name: str
    original_text: str
    improved_text: str
    explanation: str
    accepted: bool = True


@dataclass
class RewriteResult:
    """Complete rewrite result with all sections."""
    sections: List[SectionRewrite]
    overall_explanation: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class ResumeWriter:
    """
    Rewrites resume sections with improved wording.
    
    Uses AI to improve resume content while keeping the original
    structure and intent. Supports PDF and DOCX export.
    """
    
    # Common resume section headings for parsing
    SECTION_HEADINGS = [
        "summary", "professional summary", "career summary", "objective",
        "experience", "work experience", "employment history", "professional experience",
        "education", "educational background",
        "skills", "technical skills", "core competencies", "key skills",
        "projects", "personal projects", "key projects",
        "certifications", "licenses",
        "awards", "achievements", "honors",
        "publications", "research",
        "volunteer", "volunteer experience",
        "languages", "interests", "hobbies",
    ]
    
    def __init__(self, ai_analyzer=None):
        """
        Initialize ResumeWriter.
        
        Args:
            ai_analyzer: AdvancedAIAnalyzer instance for AI calls.
                        If None, creates a new one.
        """
        if ai_analyzer is None:
            from utils.ai_analyzer import AdvancedAIAnalyzer
            self.ai_analyzer = AdvancedAIAnalyzer()
        else:
            self.ai_analyzer = ai_analyzer
    
    def parse_resume_sections(self, resume_text: str) -> List[Dict[str, str]]:
        """
        Parse resume text into sections based on headings.
        
        Args:
            resume_text: Full resume text
            
        Returns:
            List of dicts with 'name' and 'content' keys
        """
        lines = resume_text.strip().split('\n')
        sections = []
        current_section = None
        current_content = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_section:
                    current_content.append('')
                continue
            
            # Check if this line is a section heading
            is_heading = False
            lower_stripped = stripped.lower().rstrip(':')
            
            # Check against known headings
            for heading in self.SECTION_HEADINGS:
                if lower_stripped == heading or lower_stripped == heading + ':':
                    is_heading = True
                    break
            
            # Also check for ALL CAPS lines (common in resumes)
            if not is_heading and stripped.isupper() and len(stripped) > 2:
                is_heading = True
            
            # Also check for lines that are short and look like headings
            if not is_heading and len(stripped) < 40 and not stripped.endswith('.') and not stripped.endswith(','):
                # Could be a heading if it's a single word or short phrase
                word_count = len(stripped.split())
                if word_count <= 4 and any(c.isupper() for c in stripped):
                    # Heuristic: if most words start with capital letters
                    capitalized = sum(1 for w in stripped.split() if w[0].isupper())
                    if capitalized / max(word_count, 1) > 0.7:
                        is_heading = True
            
            if is_heading:
                # Save previous section
                if current_section:
                    sections.append({
                        'name': current_section,
                        'content': '\n'.join(current_content).strip()
                    })
                current_section = stripped.rstrip(':')
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_section:
            sections.append({
                'name': current_section,
                'content': '\n'.join(current_content).strip()
            })
        
        # If no sections found, treat entire resume as one section
        if not sections:
            sections.append({
                'name': 'Full Resume',
                'content': resume_text.strip()
            })
        
        return sections
    
    def rewrite_resume(
        self,
        resume_text: str,
        job_description: str,
        analysis_results: Optional[Dict[str, Any]] = None
    ) -> RewriteResult:
        """
        Rewrite resume sections with improved wording.
        
        Args:
            resume_text: Original resume text
            job_description: Target job description
            analysis_results: Optional analysis results for context
            
        Returns:
            RewriteResult with rewritten sections
        """
        sections = self.parse_resume_sections(resume_text)
        
        # Extract context from analysis if available
        missing_skills = []
        priority_skills = []
        if analysis_results and 'ai_analysis' in analysis_results:
            ai_analysis = analysis_results['ai_analysis']
            if 'skill_gap_analysis' in ai_analysis:
                skill_gap = ai_analysis['skill_gap_analysis']
                missing_skills = getattr(skill_gap, 'missing_skills', []) or []
                priority_skills = getattr(skill_gap, 'priority_skills', []) or []
        
        rewritten_sections = []
        for section in sections:
            rewritten = self._rewrite_section(
                section['name'],
                section['content'],
                job_description,
                missing_skills,
                priority_skills
            )
            rewritten_sections.append(rewritten)
        
        return RewriteResult(
            sections=rewritten_sections,
            overall_explanation="Resume sections improved with better wording, "
                              "action verbs, and job-relevant keywords.",
            metadata={
                'original_section_count': len(sections),
                'missing_skills': missing_skills,
                'priority_skills': priority_skills,
            }
        )
    
    def _rewrite_section(
        self,
        section_name: str,
        content: str,
        job_description: str,
        missing_skills: List[str],
        priority_skills: List[str]
    ) -> SectionRewrite:
        """Rewrite a single section using AI."""
        
        system_prompt = """You are a professional resume writer. You MUST respond with valid JSON only — no markdown, no explanation text, no code fences.

Return a JSON object with this exact structure:
{
  "improved_text": "<the improved section text>",
  "explanation": "<brief explanation of what was changed and why>"
}

Rules:
- Keep the original structure and intent
- Improve wording with strong action verbs
- Add relevant keywords naturally (not forced)
- Quantify achievements where possible
- Keep it concise and professional
- Do NOT add information that wasn't in the original"""
        
        skills_context = ""
        if missing_skills:
            skills_context = f"\nRelevant skills to emphasize if present: {', '.join(missing_skills[:10])}"
        if priority_skills:
            skills_context += f"\nPriority skills for this role: {', '.join(priority_skills[:5])}"
        
        prompt = f"""Rewrite this resume section to better match the job description.

SECTION: {section_name}

ORIGINAL:
{content}

JOB DESCRIPTION (relevant parts):
{job_description[:1500]}
{skills_context}

Improve the wording, add relevant keywords naturally, use strong action verbs.
Respond with JSON only."""
        
        try:
            response = self.ai_analyzer._make_api_request(prompt, system_prompt)
            
            # Parse JSON response
            if "```json" in response:
                json_text = response.split("```json")[1].split("```")[0]
            else:
                json_text = response
            
            data = json.loads(json_text.strip())
            
            return SectionRewrite(
                section_name=section_name,
                original_text=content,
                improved_text=data.get('improved_text', content),
                explanation=data.get('explanation', 'Improved wording and keywords'),
                accepted=True
            )
            
        except (json.JSONDecodeError, Exception) as e:
            logger.warning(f"Failed to rewrite section '{section_name}': {e}")
            return SectionRewrite(
                section_name=section_name,
                original_text=content,
                improved_text=content,
                explanation=f"Rewrite failed: {str(e)}",
                accepted=False
            )
    
    def generate_final_version(self, sections: List[SectionRewrite]) -> str:
        """
        Generate clean final resume from accepted sections.
        
        Args:
            sections: List of SectionRewrite objects
            
        Returns:
            Clean resume text
        """
        parts = []
        for section in sections:
            if section.accepted:
                parts.append(f"{section.section_name}\n{section.improved_text}")
            else:
                parts.append(f"{section.section_name}\n{section.original_text}")
        
        return '\n\n'.join(parts)
    
    def export_pdf(self, text: str, filepath: str) -> bool:
        """
        Export resume as PDF.
        
        Args:
            text: Resume text
            filepath: Output file path
            
        Returns:
            True if successful
        """
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Parse sections from text
            lines = text.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    pdf.ln(5)
                    continue
                
                # Check if it's a section heading (short, likely capitalized)
                is_heading = (
                    len(stripped) < 40 and 
                    not stripped.endswith('.') and
                    not stripped.endswith(',') and
                    stripped.upper() == stripped
                )
                
                if is_heading:
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.cell(0, 8, stripped, new_x="LMARGIN", new_y="NEXT")
                    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                    pdf.ln(3)
                elif stripped.startswith("•") or stripped.startswith("-") or stripped.startswith("–"):
                    pdf.set_font("Helvetica", "", 10)
                    pdf.cell(5)
                    pdf.multi_cell(0, 6, stripped)
                else:
                    pdf.set_font("Helvetica", "", 10)
                    pdf.multi_cell(0, 6, stripped)
            
            pdf.output(filepath)
            logger.info(f"PDF exported to {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return False
    
    def export_docx(self, text: str, filepath: str) -> bool:
        """
        Export resume as DOCX.
        
        Args:
            text: Resume text
            filepath: Output file path
            
        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            lines = text.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                
                # Check if it's a section heading
                is_heading = (
                    len(stripped) < 40 and
                    not stripped.endswith('.') and
                    not stripped.endswith(',') and
                    stripped.upper() == stripped
                )
                
                if is_heading:
                    heading = doc.add_heading(stripped, level=2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif stripped.startswith("•") or stripped.startswith("-") or stripped.startswith("–"):
                    # Bullet point
                    p = doc.add_paragraph(stripped[1:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(stripped)
            
            doc.save(filepath)
            logger.info(f"DOCX exported to {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return False
"""
Text Extractor Module
====================

Handles extraction of text content from various file formats including PDF and TXT files.
Supports multiple PDF libraries for robust text extraction.

Copyright (c) 2025 SyazWak
Licensed under the MIT License - see LICENSE file for details.
"""

import os
import logging
from typing import Optional, Union
from pathlib import Path

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    
try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """
    A robust text extractor that handles multiple file formats.
    
    Supports:
    - PDF files (using PyMuPDF or pdfminer as fallback)
    - TXT files
    - Direct text input
    """
    
    def __init__(self, preferred_pdf_engine: str = "pymupdf"):
        """
        Initialize the TextExtractor.
        
        Args:
            preferred_pdf_engine (str): Preferred PDF processing engine ("pymupdf" or "pdfminer")
        """
        self.preferred_pdf_engine = preferred_pdf_engine
        self.supported_formats = ['.pdf', '.txt', '.docx']
        
        # Check available libraries
        if not PYMUPDF_AVAILABLE and not PDFMINER_AVAILABLE:
            logger.warning("No PDF processing libraries available. Install PyMuPDF or pdfminer.six")
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX files will not be supported.")
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a file.
        
        Args:
            file_path (Union[str, Path]): Path to the file to extract text from
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file format is not supported
            Exception: If text extraction fails
        """
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        file_extension = file_path.suffix.lower()
        
        # Check if format is supported
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file using available libraries.
        
        Args:
            file_path (Path): Path to PDF file
            
        Returns:
            str: Extracted text
            
        Raises:
            Exception: If all PDF extraction methods fail
        """
        text = ""
        
        # Try preferred engine first
        if self.preferred_pdf_engine == "pymupdf" and PYMUPDF_AVAILABLE:
            try:
                text = self._extract_with_pymupdf(file_path)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Fallback to pdfminer if PyMuPDF fails or not preferred
        if PDFMINER_AVAILABLE:
            try:
                text = self._extract_with_pdfminer(file_path)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"pdfminer extraction failed: {e}")
        
        # Try PyMuPDF as fallback if it wasn't the preferred engine
        if self.preferred_pdf_engine != "pymupdf" and PYMUPDF_AVAILABLE:
            try:
                text = self._extract_with_pymupdf(file_path)
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF fallback extraction failed: {e}")
        
        raise Exception("All PDF extraction methods failed")
    
    def _extract_with_pymupdf(self, file_path: Path) -> str:
        """Extract text using PyMuPDF (fitz)."""
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text()
        
        doc.close()
        return text
    
    def _extract_with_pdfminer(self, file_path: Path) -> str:
        """Extract text using pdfminer."""
        return pdfminer_extract(str(file_path))
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path (Path): Path to TXT file
            
        Returns:
            str: File content
        """
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any supported encoding")
    
    def validate_file(self, file_path: Union[str, Path]) -> tuple[bool, str]:
        """
        Validate if a file can be processed.
        
        Args:
            file_path (Union[str, Path]): Path to file
            
        Returns:
            tuple[bool, str]: (is_valid, error_message)
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return False, "File does not exist"
            
            if file_path.suffix.lower() not in self.supported_formats:
                return False, f"Unsupported format. Supported: {', '.join(self.supported_formats)}"
            
            if file_path.stat().st_size == 0:
                return False, "File is empty"
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (Path): Path to the DOCX file
            
        Returns:
            str: Extracted text content
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise Exception("python-docx library is not available. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                logger.warning(f"No text content found in DOCX file: {file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX file")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX file {file_path}: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def get_supported_formats(self) -> list[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()
    
    def get_available_pdf_engines(self) -> list[str]:
        """Get list of available PDF processing engines."""
        engines = []
        if PYMUPDF_AVAILABLE:
            engines.append("pymupdf")
        if PDFMINER_AVAILABLE:
            engines.append("pdfminer")
        return engines
